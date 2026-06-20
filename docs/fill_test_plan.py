from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\Tea\Downloads\TestPlan_IzvestajTestiranja.docx")
OUTPUT = Path(__file__).with_name("TestPlan_IzvestajTestiranja_Popunjen.docx")


CLASSES = [
    ("Ime je uneto i ima 1–255 karaktera.", "LK"),
    ("Ime je prazno.", "NK"),
    ("Email ima validan format, npr. korisnik@domen.rs.", "LK"),
    ("Email je prazan.", "NK"),
    ("Email nema validan format (nedostaje @ ili domen).", "NK"),
    ("Email već postoji u tabeli users.", "NK"),
    ("Lozinka ima najmanje 8 karaktera (granice: 8 i 9).", "LK"),
    ("Lozinka ima 0–7 karaktera (granična vrednost: 7).", "NK"),
    ("Potvrda lozinke je jednaka lozinci.", "LK"),
    ("Potvrda lozinke je prazna ili se ne poklapa.", "NK"),
    ("Postoji korisnik sa unetim email-om i tačnom lozinkom.", "LK"),
    ("Postoji korisnik, ali je lozinka pogrešna.", "NK"),
    ("Korisnik sa unetim email-om ne postoji.", "NK"),
    ("Datum treninga je današnji ili budući.", "LK"),
    ("Datum treninga je u prošlosti.", "NK"),
    ("Datum treninga je prazan ili nije validan datum.", "NK"),
    ("Vreme je uneto u validnom formatu HH:mm.", "LK"),
    ("Vreme je prazno.", "NK"),
    ("Vreme nije u formatu HH:mm.", "NK"),
    ("Lokacija je uneta i ima najviše 100 karaktera.", "LK"),
    ("Lokacija je prazna.", "NK"),
    ("Lokacija ima više od 100 karaktera.", "NK"),
    ("Selekcija je izabrana i postoji u bazi.", "LK"),
    ("Selekcija nije izabrana.", "NK"),
    ("Prosleđeni ID selekcije ne postoji.", "NK"),
    ("Trener postoji i dodeljen je izabranoj selekciji.", "LK"),
    ("Trener nije izabran.", "NK"),
    ("Prosleđeni ID trenera ne postoji.", "NK"),
    ("Trener postoji, ali pripada drugoj selekciji.", "NK"),
    ("Korisnik roditelj pokušava kreiranje/izmenu/brisanje treninga.", "NK"),
]


TESTS = [
    {
        "name": "Uspešna registracija – lozinka 8 karaktera",
        "description": "Registracija sa validnim imenom, email-om i donjom graničnom vrednošću lozinke.\nLK: K1, K3, K7, K9",
        "steps": "1. Otvoriti Register.\n2. Uneti ime i jedinstven validan email.\n3. Uneti lozinku „Test1234“ i istu potvrdu.\n4. Poslati formu / POST /api/register.\n5. Proveriti poruku, token i korisnika.",
        "conditions": "Preduslov: email ne postoji.\nPostuslov: kreiran user i Sanctum token; korisnik je prijavljen.",
    },
    {
        "name": "Registracija bez email-a",
        "description": "Provera obaveznog polja email.\nLK: K1, K7, K9\nNK: K4",
        "steps": "1. Popuniti ime i obe lozinke.\n2. Ostaviti email prazan.\n3. Poslati formu.\n4. Proveriti HTTP 422 i poruku „Email je obavezan“.",
        "conditions": "Preduslov: otvorena Register forma.\nPostuslov: korisnik nije kreiran.",
    },
    {
        "name": "Registracija sa nevalidnim email-om",
        "description": "Email bez znaka @ ili bez domena mora biti odbijen.\nLK: K1, K7, K9\nNK: K5",
        "steps": "1. Uneti email „tea.test“.\n2. Uneti validne ostale podatke.\n3. Poslati formu.\n4. Proveriti frontend poruku i API HTTP 422.",
        "conditions": "Preduslov: nema.\nPostuslov: korisnik nije kreiran.",
    },
    {
        "name": "Registracija – lozinka 7 karaktera",
        "description": "Provera donje nevalidne granične vrednosti.\nLK: K1, K3, K9\nNK: K8",
        "steps": "1. Uneti validno ime i email.\n2. Uneti lozinku od tačno 7 karaktera.\n3. Ponoviti istu lozinku.\n4. Poslati formu i proveriti grešku minimum 8.",
        "conditions": "Preduslov: email je jedinstven.\nPostuslov: korisnik nije kreiran.",
    },
    {
        "name": "Registracija – potvrda se ne poklapa",
        "description": "Lozinka i potvrda imaju različite vrednosti.\nLK: K1, K3, K7\nNK: K10",
        "steps": "1. Uneti validne podatke.\n2. Uneti „Test1234“.\n3. U potvrdu uneti „Test5678“.\n4. Proveriti poruku da se potvrda ne poklapa.",
        "conditions": "Preduslov: email je jedinstven.\nPostuslov: korisnik nije kreiran.",
    },
    {
        "name": "Registracija – lozinka 9 karaktera",
        "description": "Provera prve vrednosti iznad granice minimuma.\nLK: K1, K3, K7, K9",
        "steps": "1. Uneti jedinstven email.\n2. Uneti lozinku od tačno 9 karaktera i potvrdu.\n3. Poslati formu.\n4. Proveriti HTTP 201.",
        "conditions": "Preduslov: email ne postoji.\nPostuslov: korisnik i token su kreirani.",
    },
    {
        "name": "Registracija sa postojećim email-om",
        "description": "Jedinstvenost email adrese u bazi.\nLK: K1, K3, K7, K9\nNK: K6",
        "steps": "1. Uneti email postojećeg korisnika.\n2. Popuniti ostala polja validno.\n3. Poslati formu.\n4. Proveriti HTTP 422 i unique poruku.",
        "conditions": "Preduslov: korisnik sa email-om postoji.\nPostuslov: nije kreiran duplikat.",
    },
    {
        "name": "Uspešan login",
        "description": "Prijava postojećeg korisnika validnim kredencijalima.\nLK: K3, K11",
        "steps": "1. Otvoriti Angular Login na portu 4200.\n2. Uneti trener@skola.rs.\n3. Uneti tačnu lozinku.\n4. Poslati formu / POST /api/login.\n5. Proveriti Sanctum token, Laravel web sesiju i preusmerenje na PHP dashboard.",
        "conditions": "Preduslov: korisnik postoji.\nPostuslov: token i web sesija su kreirani; otvoren je PHP dashboard na portu 8000.",
    },
    {
        "name": "Login sa pogrešnom lozinkom",
        "description": "Postojeći email uz netačnu lozinku.\nLK: K3\nNK: K12",
        "steps": "1. Uneti email postojećeg korisnika.\n2. Uneti pogrešnu lozinku.\n3. Poslati formu.\n4. Proveriti HTTP 422 i jasnu poruku.",
        "conditions": "Preduslov: korisnik postoji.\nPostuslov: token nije kreiran.",
    },
    {
        "name": "Login nepostojećeg korisnika",
        "description": "Email validnog formata ne postoji u bazi.\nLK: K3\nNK: K13",
        "steps": "1. Uneti nepostoji@skola.rs.\n2. Uneti lozinku.\n3. Poslati formu.\n4. Proveriti generičku grešku kredencijala.",
        "conditions": "Preduslov: email ne postoji.\nPostuslov: korisnik nije prijavljen.",
    },
    {
        "name": "Uspešno kreiranje budućeg treninga",
        "description": "Validan termin sa povezanom selekcijom i trenerom.\nLK: K14, K17, K20, K23, K26",
        "steps": "1. Prijaviti se kao trener.\n2. Sa PHP liste otvoriti Angular formu Novi trening.\n3. Uneti budući datum, vreme i realnu lokaciju.\n4. Izabrati svoju selekciju i trenera.\n5. Sačuvati / POST /api/treninzi i proveriti povratak na PHP listu.",
        "conditions": "Preduslov: trener je prijavljen i ima dodeljenu selekciju.\nPostuslov: trening je u bazi i prikazan na PHP listi treninga.",
    },
    {
        "name": "Kreiranje treninga bez lokacije",
        "description": "Obaveznost lokacije na frontend i backend nivou.\nLK: K14, K17, K23, K26\nNK: K21",
        "steps": "1. Popuniti sve osim lokacije.\n2. Poslati formu.\n3. Proveriti poruku „Lokacija je obavezna“.\n4. U Postmanu proveriti HTTP 422.",
        "conditions": "Preduslov: trener je prijavljen.\nPostuslov: trening nije kreiran.",
    },
    {
        "name": "Kreiranje treninga u prošlosti",
        "description": "Datum pre današnjeg datuma nije dozvoljen.\nLK: K17, K20, K23, K26\nNK: K15",
        "steps": "1. Uneti jučerašnji datum.\n2. Popuniti ostala polja validno.\n3. Poslati zahtev.\n4. Proveriti HTTP 422 i poruku o prošlosti.",
        "conditions": "Preduslov: trener je prijavljen.\nPostuslov: trening nije kreiran.",
    },
    {
        "name": "Kreiranje treninga za današnji datum",
        "description": "Današnji datum je validna granična vrednost.\nLK: K14, K17, K20, K23, K26",
        "steps": "1. Uneti današnji datum.\n2. Uneti validno vreme, lokaciju, selekciju i trenera.\n3. Poslati zahtev.\n4. Proveriti HTTP 201.",
        "conditions": "Preduslov: trener je prijavljen.\nPostuslov: trening je kreiran.",
    },
    {
        "name": "Kreiranje bez selekcije i trenera",
        "description": "Provera oba obavezna povezana entiteta.\nLK: K14, K17, K20\nNK: K24, K27",
        "steps": "1. Uneti datum, vreme i lokaciju.\n2. Ne izabrati selekciju ni trenera.\n3. Poslati formu.\n4. Proveriti obe validacione poruke i HTTP 422.",
        "conditions": "Preduslov: trener je prijavljen.\nPostuslov: trening nije kreiran.",
    },
    {
        "name": "Trener nije iz izabrane selekcije",
        "description": "Poslovno pravilo veze trener–selekcija.\nLK: K14, K17, K20, K23\nNK: K29",
        "steps": "1. U Postmanu izabrati validnu selekciju.\n2. Poslati ID trenera druge selekcije.\n3. Poslati POST /api/treninzi.\n4. Proveriti HTTP 422 za trener_id.",
        "conditions": "Preduslov: postoje dve selekcije i treneri.\nPostuslov: trening nije kreiran.",
    },
]


DEFECTS = [
    ("8", "Angular login je kreirao API token, ali Laravel PHP deo nije prepoznavao web sesiju zbog nepotpunog session middleware-a. API login/register su prebačeni na web middleware; retest prošao."),
    ("8", "Posle PHP logout-a Angular localStorage je kratko prikazivao prethodnog korisnika. Dodat je logout signal koji briše Angular token i korisnika; retest prošao."),
    ("11", "U demo bazi su nakon testiranja ostali članovi Browser/Postman i lokacije Angular teren. Seederi su očišćeni, dodati su realni članovi i lokacije, baza je ponovo seedovana; retest prošao."),
]


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches):
    width = Inches(inches)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width / 635)))
    tc_w.set(qn("w:type"), "dxa")


def write_cell(cell, text, *, bold=False, size=8, center=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table(table, widths, body_size=8):
    table.autofit = False
    set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(body_size if row_index else body_size)
                    if row_index == 0:
                        run.bold = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade_cell(cell, "DDEBE3")


doc = Document(SOURCE)

doc.paragraphs[10].text = "Testiranje softvera – Informacioni sistem škole fudbala"
doc.paragraphs[25].text = "Student: Tea Samčević"
doc.paragraphs[26].text = "Broj indeksa: __________________"
doc.paragraphs[27].text = "Procesi: Login/Register i zakazivanje/kreiranje treninga"

for paragraph in [doc.paragraphs[9], doc.paragraphs[29], doc.paragraphs[31], doc.paragraphs[41], doc.paragraphs[46]]:
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(21, 95, 70)

classes_table = doc.tables[0]
for index, (description, class_type) in enumerate(CLASSES, start=1):
    write_cell(classes_table.cell(index, 0), f"{index}.", center=True)
    write_cell(classes_table.cell(index, 1), description)
    write_cell(classes_table.cell(index, 2), class_type, bold=True, center=True)
format_table(classes_table, [0.43, 5.85, 0.85], body_size=8)

test_table = doc.tables[1]
for index, test in enumerate(TESTS, start=1):
    write_cell(test_table.cell(index, 0), f"{index}.", center=True, size=7)
    write_cell(test_table.cell(index, 1), test["name"], bold=True, size=7)
    write_cell(test_table.cell(index, 2), test["description"], size=7)
    write_cell(test_table.cell(index, 3), test["steps"], size=7)
    write_cell(test_table.cell(index, 4), test["conditions"], size=7)
    write_cell(test_table.cell(index, 5), "Manuelno kroz Angular UI\n+\nPostman API testiranje", size=7, center=True)
for index in range(len(TESTS) + 1, len(test_table.rows)):
    for column in range(len(test_table.columns)):
        write_cell(test_table.cell(index, column), "")
format_table(test_table, [0.34, 0.92, 1.25, 1.82, 1.46, 1.47], body_size=7)

report_table = doc.tables[2]
for index in range(1, len(report_table.rows)):
    if index <= len(TESTS):
        write_cell(report_table.cell(index, 0), f"{index}.", center=True)
        write_cell(report_table.cell(index, 1), "20.06.2026.", center=True)
        write_cell(report_table.cell(index, 2), "Izvršen (manuelno i Postman)", center=True)
        write_cell(report_table.cell(index, 3), "Prošao", bold=True, center=True)
    else:
        for column in range(4):
            write_cell(report_table.cell(index, column), "")
format_table(report_table, [0.75, 1.25, 3.15, 1.25], body_size=8)

summary_table = doc.tables[3]
summary_values = [
    ("Nije izvršeno", "0"),
    ("Izvršeno (manuelno)", str(len(TESTS))),
    ("Izvršeno (Postman)", str(len(TESTS))),
    ("Prošlo", str(len(TESTS))),
    ("Nije prošlo", "0"),
    ("Ukupno test primera", str(len(TESTS))),
]
for row_index, values in enumerate(summary_values, start=1):
    write_cell(summary_table.cell(row_index, 0), values[0], bold=row_index == 6)
    write_cell(summary_table.cell(row_index, 1), values[1], bold=True, center=True)
format_table(summary_table, [3.55, 3.55], body_size=9)

defect_table = doc.tables[4]
for index, (test_case, description) in enumerate(DEFECTS, start=1):
    write_cell(defect_table.cell(index, 0), f"{index}.", center=True)
    write_cell(defect_table.cell(index, 1), test_case, center=True)
    write_cell(defect_table.cell(index, 2), description)
for index in range(len(DEFECTS) + 1, len(defect_table.rows)):
    for column in range(3):
        write_cell(defect_table.cell(index, column), "")
format_table(defect_table, [0.5, 0.9, 5.7], body_size=8)

def add_section_heading(text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(21, 95, 70)


add_section_heading("KRATAK REZIME TESTIRANJA")
doc.add_paragraph(
    f"Izvršeno je ukupno {len(TESTS)} test primera za procese Login/Register i "
    "zakazivanje/kreiranje treninga. Svih 16 testova izvršeno je manuelno kroz "
    "Angular korisnički interfejs i kroz Postman REST API zahteve. Nakon otklanjanja "
    "tri defekta i ponovnog testiranja, 16 testova je prošlo, 0 nije prošlo, a 0 je ostalo neizvršeno."
)
add_section_heading("TESTNO OKRUŽENJE")
doc.add_paragraph(
    "Laravel 11 REST API i postojeći PHP/Blade deo rade na portu 8000, dok Angular 21 "
    "Login, Register i forma za kreiranje treninga rade na portu 4200. Komunikacija je JSON, "
    "autentifikacija koristi Laravel Sanctum token i Laravel web sesiju. SQLite je korišćen "
    "za demo i automatizovane testove, a projekat podržava MySQL razvojnu bazu."
)
add_section_heading("NAPOMENA O AUTOMATIZOVANOJ PROVERI")
doc.add_paragraph(
    "Pored manuelnog i Postman plana, pokrenut je Laravel PHPUnit paket: 8 testova i "
    "42 assertions su prošli. Postman/Newman kolekcija izvršila je 19 zahteva i 27 assertions "
    "bez greške. Angular produkcioni build je uspešno generisan. Browser proverom potvrđeni su "
    "Login, Register, trenerovo kreiranje treninga, ograničenje selekcije i povratak na PHP listu."
)

for paragraph in doc.paragraphs:
    if paragraph.text and paragraph.style.name == "Normal":
        for run in paragraph.runs:
            run.font.name = "Arial"

doc.save(OUTPUT)
print(OUTPUT)
